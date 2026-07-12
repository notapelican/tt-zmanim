"""Document generator (warplan module 4, presentation half): render the plain
block data from `engine.assemble.generate()` into .docx sheets matching the
TTCC house style (see RENDERER-CONTRACT.md and samples/).

This module performs layout/styling ONLY. It never computes or re-rounds a
zman, and never invents a time, label, or note that is not already present in
the block data -- see RENDERER-CONTRACT.md for the data contract.

Three sheet formats, chosen automatically by `render_docx()` from the shape
of the generated data:
  - single week  (1 week block)               -> one column, full width
  - multi-week   (2+ week blocks)              -> two-column layout
  - yom tov day  (DAY blocks, interleaved with their week) -> boxed
    day-by-day schedule, rendered inline where the day falls
"""
from __future__ import annotations

from datetime import date as _date

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Mm, Pt, RGBColor

# --- house style constants (colors sampled from the historical PDFs) -------
BLUE = RGBColor(0x00, 0x00, 0xFF)          # titles, section bars
WHITE = RGBColor(0xFF, 0xFF, 0xFF)         # bar text
BLACK = RGBColor(0x00, 0x00, 0x00)
FAST_FILL = "FFF2CC"                        # cream/tan box fill
FAST_BORDER_HEX = "EE0000"
BAR_FILL_BLUE = "0000FF"
BAR_FILL_PURPLE = "800080"

# Times New Roman matches the historical sheets exactly (confirmed from the
# PDFs' embedded font names); LibreOffice/Word fall back to a Hebrew-capable
# font for בס"ד within the same run automatically.
FONT_NAME = "Times New Roman"
BODY_SIZE = Pt(11)          # single-week sheet
MULTI_BODY_SIZE = Pt(9)     # multi-week two-column sheet (narrower columns)

PAGE_MARGIN = Mm(13)

# Canonical print order for week-block sections; matches the historical
# sheets, which do NOT follow the raw `entries` list order for two spots:
# the early-minyan section prints after "Erev Shabbos key times" (not
# before, as the raw data has it), and any `kind: "fast"` box prints right
# after the section-less zmanim table, before the weekday section.
_SEC_FAST = "\x00fast"
_SEC_NONE = None
WEEKDAY = "Davening times during the week"
EARLY_ES = "Erev Shabbos early minyan (and essential halachic details)"
KEY_TIMES = "Erev Shabbos key times"
SHABBOS_DAY = "Shabbos Day and Motzaei Shabbos"

_SECTION_PRIORITY = {
    _SEC_NONE: 0,
    _SEC_FAST: 1,
    WEEKDAY: 2,
    KEY_TIMES: 3,
    EARLY_ES: 4,
    # Erev Shabbos candle-lighting section (title varies) is anything not
    # otherwise matched between EARLY_ES and SHABBOS_DAY -- handled below.
    SHABBOS_DAY: 6,
}
_EREV_SHABBOS_SLOT = 5

# Within SHABBOS_DAY, the raw entries list order does not match print order
# either (Tehillim/Chassidus prints before Morning Shema; Halacha shiur
# prints before Mincha/Motzaei). Recovered from the historical sheets.
_SHABBOS_DAY_RULE_PRIORITY = {
    "shab_tehillim": 0, "shab_chassidus": 0,
    "z_shema_shab": 1,
    "shab_shacharis_mev": 2, "shab_shacharis": 2,
    # molad paragraph (not a LINE) is inserted right after slot 2
    "shab_halacha_shiur": 3,
    "shab_mincha": 4,
    "motzaei_maariv": 5,
}


def _fmt_ampm(hhmm: str) -> str:
    h, m = int(hhmm[:2]), int(hhmm[3:5])
    ampm = "am" if h < 12 else "pm"
    h12 = h % 12 or 12
    return f"{h12}:{m:02d}{ampm}"


# House month style: short names in full (May, June, July), the rest
# abbreviated with a period (Oct., Sept., ...) as printed on the sheets.
_MONTH_STYLE = {"May": "May", "Jun": "June", "Jul": "July", "Sep": "Sept."}


def _mon(d: _date) -> str:
    m = d.strftime("%b")
    return _MONTH_STYLE.get(m, m + ".")


def _fmt_civil_range(start_iso: str, end_iso: str) -> str:
    s, e = _date.fromisoformat(start_iso), _date.fromisoformat(end_iso)
    if s.month == e.month:
        return f"{s.day}–{e.day} {_mon(e)} ’{e:%y}"
    return f"{s.day} {_mon(s)} – {e.day} {_mon(e)} ’{e:%y}"


def _fmt_civil_date(iso: str) -> str:
    d = _date.fromisoformat(iso)
    return f"{d.day} {_mon(d)}"


# --- low-level docx helpers -------------------------------------------------

def _set_cell_background(cell, hex_color: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_borders(cell, **sides) -> None:
    """sides: e.g. right={'sz':12,'color':'000000','val':'single'}"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side, spec in sides.items():
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), spec.get("val", "nil"))
        el.set(qn("w:sz"), str(spec.get("sz", 0)))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), spec.get("color", "auto"))
        borders.append(el)
    tcPr.append(borders)


def _no_borders(cell) -> None:
    _set_cell_borders(cell, top={}, bottom={}, left={}, right={})


def _set_cell_margins(cell, twips=60) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(twips))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _set_col_widths(table, widths) -> None:
    widths = [Emu(int(w)) for w in widths]
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths):
            cell.width = w
    grid = table._tbl.find(qn("w:tblGrid"))
    for gridCol, w in zip(grid.findall(qn("w:gridCol")), widths):
        gridCol.set(qn("w:w"), str(w.twips))


def _para(container, text="", *, bold=False, italic=False, size=BODY_SIZE,
          color=BLACK, align=None, space_after=Pt(2), space_before=Pt(0)):
    p = container.add_paragraph()
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = space_before
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.font.name = FONT_NAME
        r.font.size = size
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = color
    return p


def _shade_run(r, fill):
    rPr = r._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    rPr.append(shd)


def _bar(container, text, width, *, purple=False, size=BODY_SIZE):
    """A colored highlight bar hugging its text (section header), matching
    the blue/purple highlighted headings in the historical sheets."""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = size
    r.font.bold = True
    r.font.color.rgb = WHITE
    _shade_run(r, BAR_FILL_PURPLE if purple else BAR_FILL_BLUE)


def _dotted_line(container, label, value, width, *, size=BODY_SIZE, bullet=False):
    p = container.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.tab_stops.add_tab_stop(
        width, WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
    prefix = "\u2022 " if bullet else ""
    r = p.add_run(f"{prefix}{label} \t {value}")
    r.font.name = FONT_NAME
    r.font.size = size
    return p


def _fast_box(container, text, width, *, size=BODY_SIZE):
    table = container.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    width = Emu(int(width * 0.9))
    cell = table.rows[0].cells[0]
    _set_cell_background(cell, FAST_FILL)
    _set_cell_borders(cell, **{
        s: {"val": "single", "sz": 12, "color": FAST_BORDER_HEX}
        for s in ("top", "bottom", "left", "right")})
    _set_cell_margins(cell, 100)
    _set_col_widths(table, [width])
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = FONT_NAME
    r.font.size = size
    r.font.bold = True
    r.font.italic = True
    _para(container, "", space_after=Pt(2))


# --- entry (LINE) merge/formatting logic ------------------------------------

def _qualified(qualifier, time_s):
    t = _fmt_ampm(time_s)
    return f"{qualifier} {t}" if qualifier else t


def _join_dayspec_group(entries):
    """Join same-(section,label,day_spec) entries: multiple qualifiers join
    with ', ' unless the next qualifier starts with 'or' (already reads as a
    continuation, e.g. 'before 8:00pm or after 8:30pm'); multiple same
    (typically qualifier-less) entries join with ' & ' (e.g. Shacharis
    8:00am & 9:15am)."""
    parts = []
    for e in entries:
        piece = _qualified(e["qualifier"], e["time"])
        if not parts:
            parts.append(piece)
        elif e["qualifier"] and e["qualifier"].startswith("or"):
            parts[-1] = f"{parts[-1]} {piece}"
        elif e["qualifier"]:
            parts[-1] = f"{parts[-1]}, {piece}"
        else:
            parts[-1] = f"{parts[-1]} & {piece}"
    return "".join(parts)


def _render_label_run(container, run, width, *, dayspec_before_leader, size=BODY_SIZE, bullet=False):
    """Render one run of consecutive entries sharing (section, label)."""
    label = run[0]["label"]
    # sub-group by day_spec, preserving first-seen order
    groups: list[tuple[str | None, list[dict]]] = []
    for e in run:
        if groups and groups[-1][0] == e["day_spec"]:
            groups[-1][1].append(e)
        else:
            groups.append((e["day_spec"], [e]))

    if len(groups) == 1:
        day_spec, group_entries = groups[0]
        value = _join_dayspec_group(group_entries)
        if day_spec and dayspec_before_leader:
            _dotted_line(container, f"{label} {day_spec}", value, width, size=size, bullet=bullet)
        elif day_spec:
            _dotted_line(container, label, f"{day_spec} {value}", width, size=size, bullet=bullet)
        else:
            _dotted_line(container, label, value, width, size=size, bullet=bullet)
        return

    # Multiple day_spec groups for the same label -- e.g. "Shacharis: Sun.
    # 8:00am & 9:15am; Mon.-Fri. 6:15am & 7:30am" (contract's own example
    # uses a colon here; other multi-group labels like Mincha/Maariv do not,
    # per the historical sheets).
    joined = "; ".join(
        f"{day_spec} {_join_dayspec_group(group_entries)}" if day_spec
        else _join_dayspec_group(group_entries)
        for day_spec, group_entries in groups)
    label_text = f"{label}:" if label == "Shacharis" else label
    _dotted_line(container, label_text, joined, width, size=size, bullet=bullet)


def _render_entry_group(container, entries, width, *, dayspec_before_leader=False, size=BODY_SIZE, bullet=False):
    """Render a same-section run of entries: merge same (label, day_spec)
    qualifier variants, then merge consecutive same-label runs across
    day_specs (see RENDERER-CONTRACT.md merge rules)."""
    # rule A: merge (label, day_spec) qualifier variants first, in order
    merged: list[dict] = []
    for e in entries:
        if (merged and merged[-1]["label"] == e["label"]
                and merged[-1]["day_spec"] == e["day_spec"]):
            merged[-1]["_group"].append(e)
        else:
            merged.append(dict(e, _group=[e]))

    # rule B: run of consecutive same-label entries (any day_spec)
    i = 0
    while i < len(merged):
        j = i
        while j + 1 < len(merged) and merged[j + 1]["label"] == merged[i]["label"]:
            j += 1
        run = [g for m in merged[i:j + 1] for g in m["_group"]]
        _render_label_run(container, run, width, dayspec_before_leader=dayspec_before_leader, size=size, bullet=bullet)
        i = j + 1


def _fast_section_title(entries) -> str:
    return entries[0]["section"]


def _partition_week_entries(entries):
    """Group entries by section, in the canonical print order (see
    _SECTION_PRIORITY), splitting out kind=='fast' runs as their own boxes."""
    zmanim_table: list[dict] = []
    fast_runs: list[list[dict]] = []
    named: dict[str, list[dict]] = {}
    order: list[str] = []  # first-seen order of named (non-fixed) sections

    i = 0
    while i < len(entries):
        e = entries[i]
        if e["kind"] == "fast":
            j = i
            run = []
            while j < len(entries) and entries[j]["kind"] == "fast" and entries[j]["section"] == e["section"]:
                run.append(entries[j])
                j += 1
            fast_runs.append(run)
            i = j
            continue
        if e["section"] is None:
            zmanim_table.append(e)
        else:
            named.setdefault(e["section"], []).append(e)
            if e["section"] not in order:
                order.append(e["section"])
        i += 1

    def slot(section: str) -> int:
        if section in _SECTION_PRIORITY:
            return _SECTION_PRIORITY[section]
        if section == KEY_TIMES:
            return _SECTION_PRIORITY[KEY_TIMES]
        return _EREV_SHABBOS_SLOT  # the (possibly renamed) Erev Shabbos section

    named_order = sorted(order, key=slot)
    return zmanim_table, fast_runs, named, named_order


def render_week_into(container, block: dict, width, *, size=BODY_SIZE,
                     notes_inline=False) -> None:
    title_size = Pt(size.pt + 2.5)
    _para(container, block["title"], bold=True, color=BLUE, size=title_size,
          space_after=Pt(0))
    _para(container, f"{block['hebrew_dates']}  "
          f"({_fmt_civil_range(block['civil_start'], block['civil_end'])})",
          bold=True, color=BLUE, size=title_size, space_after=Pt(6))

    zmanim_table, fast_runs, named, named_order = _partition_week_entries(block["entries"])

    for e in zmanim_table:
        _render_label_run(container, [e], width, dayspec_before_leader=True, size=size)

    for run in fast_runs:
        start, end = run[0], run[1]
        title = _fast_section_title(run)
        # Same-day fast prints "(Thurs.) starts at ...; ends at ..." per the
        # sheets; a fast spanning days names each day.
        if start["day_spec"] == end["day_spec"]:
            text = (f"{title} ({start['day_spec']}): starts at "
                    f"{_fmt_ampm(start['time'])}; ends at {_fmt_ampm(end['time'])}.")
        else:
            text = (f"{title}: starts {start['day_spec']} at {_fmt_ampm(start['time'])}; "
                    f"ends {end['day_spec']} at {_fmt_ampm(end['time'])}.")
        _fast_box(container, text, width, size=size)

    if notes_inline and block.get("notes"):
        for n in block["notes"]:
            _para(container, n, italic=True, size=size)

    shabbos_bar_done = False

    def emit_shabbos_bar():
        nonlocal shabbos_bar_done
        if shabbos_bar_done:
            return
        labels = ", ".join([block["parsha"]] + block["shabbos_labels"])
        _bar(container, f"Shabbos kodesh: {labels}", width, purple=True, size=size)
        shabbos_bar_done = True

    for section in named_order:
        section_entries = named[section]
        if section == WEEKDAY:
            _bar(container, section, width, size=size)
            _render_entry_group(container, section_entries, width, size=size)
            continue
        # everything Friday/Shabbos-related is introduced by the "Shabbos
        # kodesh:" bar, inserted right after the weekday section (matches
        # the historical sheets; the raw entries order alone does not).
        emit_shabbos_bar()
        if section == KEY_TIMES:
            _para(container, section, bold=True, space_after=Pt(2), size=size)
            _render_entry_group(container, section_entries, width, size=size)
        elif section == SHABBOS_DAY:
            _bar(container, section, width, size=size)
            _render_shabbos_day(container, section_entries, width, block.get("molad"), size=size)
        else:
            _bar(container, section, width, size=size)
            _render_entry_group(container, section_entries, width, size=size,
                                bullet=(section == EARLY_ES))

    emit_shabbos_bar()  # in case a week has no Erev Shabbos content at all

    if block.get("notes") and not notes_inline:
        _notes_foot(container, block["notes"], width, size=size)


def _render_shabbos_day(container, entries, width, molad, *, size=BODY_SIZE):
    ordered = sorted(enumerate(entries),
                      key=lambda ie: (_SHABBOS_DAY_RULE_PRIORITY.get(ie[1]["rule_id"], 99), ie[0]))
    ordered_entries = [e for _, e in ordered]
    molad_size = Pt(size.pt - 1)
    if not molad:
        _render_entry_group(container, ordered_entries, width, size=size)
        return
    split = None
    for idx, e in enumerate(ordered_entries):
        if e["rule_id"] in ("shab_shacharis_mev", "shab_shacharis"):
            split = idx + 1
            break
    if split is None:
        _render_entry_group(container, ordered_entries, width, size=size)
        _para(container, molad, italic=True, size=molad_size)
        return
    _render_entry_group(container, ordered_entries[:split], width, size=size)
    mp = _para(container, molad, italic=True, size=molad_size)
    mp.paragraph_format.left_indent = Mm(3)
    _render_entry_group(container, ordered_entries[split:], width, size=size)


def _notes_foot(container, notes, width, *, size=BODY_SIZE):
    note_size = Pt(size.pt - 1)
    table = container.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _no_borders(cell)
    _set_cell_borders(cell, top={"val": "single", "sz": 4, "color": "000000"})
    _set_cell_margins(cell, 40)
    _set_col_widths(table, [width])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    for i, note in enumerate(notes):
        if i:
            p = cell.add_paragraph()
        r = p.add_run(note)
        r.font.name = FONT_NAME
        r.font.size = note_size
    _para(container, " ", space_after=Pt(0), size=Pt(4))


def render_day_into(container, block: dict, width, *, size=BODY_SIZE) -> None:
    title = block["title"] or ", ".join(block["labels"])
    heading = f"{title}: {block['hebrew_date']} ({block['weekday']} {_fmt_civil_date(block['date'])})"
    _bar(container, heading, width, size=size)
    if block.get("omer_day"):
        _para(container, f"Day {block['omer_day']} of the Omer", italic=True, size=size)
    _render_entry_group(container, block["entries"], width, size=size)
    if block.get("notes"):
        _notes_foot(container, block["notes"], width, size=size)


# --- top-level assembly -----------------------------------------------------

def _group_blocks(blocks: list[dict]) -> list[dict]:
    groups = []
    cur = None
    for b in blocks:
        if b["type"] == "week":
            cur = {"week": b, "days": []}
            groups.append(cur)
        else:
            cur["days"].append(b)
    return groups


def _render_group(container, group, width, *, size=BODY_SIZE, notes_inline=False):
    render_week_into(container, group["week"], width, size=size,
                     notes_inline=notes_inline)
    for day in group["days"]:
        render_day_into(container, day, width, size=size)
        _para(container, " ", space_after=Pt(0), size=Pt(4))


def _week_separator(container):
    """Thin horizontal rule between stacked week blocks in a column."""
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def _add_header(doc, width, *, multi=False):
    table = doc.add_table(rows=1, cols=2)
    _no_borders(table.rows[0].cells[0])
    _no_borders(table.rows[0].cells[1])
    _set_col_widths(table, [width * 0.8, width * 0.2])
    title_cell, bsd_cell = table.rows[0].cells
    p = title_cell.paragraphs[0]
    r = p.add_run("Tzemach Tzedek Community Centre")
    r.font.name = FONT_NAME
    r.font.size = Pt(18)
    r.font.bold = True
    r.font.color.rgb = BLUE
    if multi:
        r2 = p.add_run("        www.ttcc.info")
        r2.font.name = FONT_NAME
        r2.font.size = Pt(13)
        r2.font.bold = True
        r2.font.color.rgb = BLUE
        r2.font.underline = True
    p2 = bsd_cell.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("בס״ד")  # בס"ד (gershayim, not an ASCII quote)
    r2.font.size = Pt(14)
    r2.font.bold = True
    if multi:
        # compact one-line address, as on the multi-week sheets
        _para(doc, "Location: 1 Penkivil St, Bondi, NSW.  "
              "Mailing address: PO Box 477 Waverley NSW 2024",
              bold=True, color=BLUE, size=Pt(12), space_after=Pt(2))
    else:
        _para(doc, "1 Penkivil St, Bondi, NSW.    www.ttcc.info", bold=True, color=BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER)
        _para(doc, "Mailing address: PO Box 477 Waverley NSW 2024", bold=True, color=BLUE,
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
    hr = doc.add_paragraph()
    hr.paragraph_format.space_after = Pt(6)
    pPr = hr._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "double")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    pbdr.append(bottom)
    pPr.append(pbdr)


def render_docx(doc_data: dict) -> Document:
    """Render `engine.assemble.generate()` output to a python-docx Document.
    Chooses single-week / multi-week-two-column / yom-tov-day layout from
    the shape of the data; day blocks always render inline where they fall
    chronologically, in whichever column their week occupies."""
    doc = Document()
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = section.right_margin = PAGE_MARGIN
    section.top_margin = section.bottom_margin = PAGE_MARGIN
    usable_width = section.page_width - section.left_margin - section.right_margin

    groups = _group_blocks(doc_data["blocks"])
    _add_header(doc, usable_width, multi=len(groups) > 1)

    shared_notes: list[str] = []
    if len(groups) > 1:
        common = set.intersection(*(set(g["week"]["notes"]) for g in groups))
        shared_notes = [n for n in groups[0]["week"]["notes"] if n in common]
        for g in groups:
            g["week"] = dict(g["week"],
                             notes=[n for n in g["week"]["notes"] if n not in common])

    if len(groups) <= 1:
        for g in groups:
            _render_group(doc, g, usable_width)
        return doc

    # multi-week two-column layout: first half of weeks (chronological) in
    # the left column, remainder in the right column -- matches the
    # historical multi-week sheets' reading order.
    half = (len(groups) + 1) // 2
    col_width = Emu(int((usable_width - Mm(4)) / 2))
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_col_widths(table, [col_width, col_width])
    left_cell, right_cell = table.rows[0].cells
    _no_borders(left_cell)
    _no_borders(right_cell)
    _set_cell_borders(left_cell, right={"val": "single", "sz": 6, "color": "000000"})
    _set_cell_margins(left_cell, 80)
    _set_cell_margins(right_cell, 120)

    for i, g in enumerate(groups[:half]):
        if i:
            _week_separator(left_cell)
        _render_group(left_cell, g, col_width, size=MULTI_BODY_SIZE, notes_inline=True)
    for i, g in enumerate(groups[half:]):
        if i:
            _week_separator(right_cell)
        _render_group(right_cell, g, col_width, size=MULTI_BODY_SIZE, notes_inline=True)
    for n in shared_notes:
        _para(doc, f"Note: {n}", size=MULTI_BODY_SIZE, space_before=Pt(6),
              align=WD_ALIGN_PARAGRAPH.CENTER)
    return doc


def save_docx(doc_data: dict, path: str) -> None:
    render_docx(doc_data).save(path)
